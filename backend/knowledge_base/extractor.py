import asyncio
import json
import os
from pathlib import Path

from knowledge_base.chunker import chunk_text


def _extract_pdf(file_path: str) -> str:
    import fitz  # PyMuPDF
    doc = fitz.open(file_path)
    pages = [page.get_text("text") for page in doc]
    doc.close()
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    import docx
    document = docx.Document(file_path)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n\n".join(parts)


async def extract_text(file_path: str, file_type: str) -> str:
    loop = asyncio.get_running_loop()
    if file_type == "pdf":
        return await loop.run_in_executor(None, _extract_pdf, file_path)
    elif file_type == "docx":
        return await loop.run_in_executor(None, _extract_docx, file_path)
    raise ValueError(f"Unsupported file type: {file_type}")


def get_chunk_sidecar_path(file_path: str) -> str:
    return file_path + ".chunks.json"


def save_chunks(file_path: str, chunks: list[str]) -> None:
    sidecar = get_chunk_sidecar_path(file_path)
    with open(sidecar, "w", encoding="utf-8") as f:
        json.dump(chunks, f, ensure_ascii=False)


def load_chunks(file_path: str) -> list[str]:
    sidecar = get_chunk_sidecar_path(file_path)
    if not os.path.exists(sidecar):
        return []
    with open(sidecar, encoding="utf-8") as f:
        return json.load(f)


async def extract_and_chunk(file_path: str, file_type: str) -> list[str]:
    text = await extract_text(file_path, file_type)
    chunks = chunk_text(text)
    save_chunks(file_path, chunks)
    return chunks
